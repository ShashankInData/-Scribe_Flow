from datetime import timedelta
from typing import Any, Dict, List, Union
import io

import srt
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


TranscriptionInput = Union[Dict[str, Any], str]


def _compose_plain_text(transcription_data: TranscriptionInput) -> str:
    """
    Build a readable plain-text transcript from either:
    - dict with "segments" (optional "speaker"), or
    - raw string
    """
    if isinstance(transcription_data, dict) and "segments" in transcription_data:
        lines: List[str] = []
        for seg in transcription_data["segments"]:
            speaker = seg.get("speaker")
            text = (seg.get("text") or "").strip()
            if not text:
                continue
            if speaker:
                lines.append(f"{speaker}: {text}")
            else:
                lines.append(text)
        return "\n\n".join(lines)
    # fallback: string content
    return str(transcription_data)


class ExportManager:
    def __init__(self):
        self.styles = getSampleStyleSheet()

    # ------------------------
    # Subtitles: SRT
    # ------------------------
    def to_srt(self, transcription_data: TranscriptionInput) -> str:
        """
        Convert transcription to SRT.
        - If dict with segments: use real timings.
        - If plain text: fallback to naive 3s-per-sentence.
        """
        if isinstance(transcription_data, dict) and "segments" in transcription_data:
            subs: List[srt.Subtitle] = []
            for i, seg in enumerate(transcription_data["segments"]):
                start = float(seg.get("start", 0.0))
                end = float(seg.get("end", max(start + 0.5, start)))
                speaker = seg.get("speaker") or ""
                text = (seg.get("text") or "").strip()
                if not text:
                    continue
                content = f"{speaker}: {text}" if speaker else text
                subs.append(
                    srt.Subtitle(
                        index=i + 1,
                        start=timedelta(seconds=start),
                        end=timedelta(seconds=end),
                        content=content,
                    )
                )
            return srt.compose(subs)
        else:
            return self._text_to_srt_fallback(str(transcription_data))

    # ------------------------
    # Subtitles: WebVTT
    # ------------------------
    def to_vtt(self, transcription_data: TranscriptionInput) -> str:
        """
        Convert transcription to WebVTT.
        """
        def _fmt_vtt_ts(seconds: float) -> str:
            # WebVTT allows hours, e.g. "00:01:02.345"
            ms = int(round((seconds - int(seconds)) * 1000))
            total = int(seconds)
            h = total // 3600
            m = (total % 3600) // 60
            s = total % 60
            return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

        if isinstance(transcription_data, dict) and "segments" in transcription_data:
            out = ["WEBVTT", ""]
            for seg in transcription_data["segments"]:
                start = float(seg.get("start", 0.0))
                end = float(seg.get("end", max(start + 0.5, start)))
                speaker = seg.get("speaker") or ""
                text = (seg.get("text") or "").strip()
                if not text:
                    continue
                content = f"{speaker}: {text}" if speaker else text
                out.append(f"{_fmt_vtt_ts(start)} --> {_fmt_vtt_ts(end)}")
                out.append(content)
                out.append("")  # blank line
            return "\n".join(out)
        else:
            return self._text_to_vtt_fallback(str(transcription_data))

    # ------------------------
    # Documents: DOCX
    # ------------------------
    def to_docx(self, transcription_data: TranscriptionInput) -> bytes:
        """
        Build a .docx file (bytes).
        """
        doc = Document()
        doc.add_heading("Transcription", 0)

        if isinstance(transcription_data, dict) and "segments" in transcription_data:
            for seg in transcription_data["segments"]:
                speaker = seg.get("speaker")
                text = (seg.get("text") or "").strip()
                if not text:
                    continue
                if speaker:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{speaker}: ")
                    run.bold = True
                    p.add_run(text)
                else:
                    doc.add_paragraph(text)
        else:
            doc.add_paragraph(str(transcription_data))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ------------------------
    # Documents: PDF
    # ------------------------
    def to_pdf(self, transcription_data: TranscriptionInput) -> bytes:
        """
        Build a .pdf file (bytes) using reportlab.
        """
        content = _compose_plain_text(transcription_data)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        story: List[Any] = []

        story.append(Paragraph("Transcription", self.styles["Title"]))
        story.append(Spacer(1, 12))

        # Split content into paragraphs at blank lines for nicer layout
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            story.append(Paragraph(block, self.styles["Normal"]))
            story.append(Spacer(1, 8))

        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    # ------------------------
    # Fallbacks for plain-text
    # ------------------------
    def _text_to_srt_fallback(self, text: str) -> str:
        """
        Naive SRT: split by '. ' and assign 3s per sentence.
        """
        parts = [p.strip() for p in text.split(". ") if p.strip()]
        subs: List[srt.Subtitle] = []
        for i, sentence in enumerate(parts):
            start = timedelta(seconds=i * 3)
            end = timedelta(seconds=(i + 1) * 3)
            subs.append(srt.Subtitle(index=i + 1, start=start, end=end, content=sentence))
        return srt.compose(subs)

    def _text_to_vtt_fallback(self, text: str) -> str:
        """
        Naive VTT: split by '. ' and assign 3s per sentence.
        """
        def _fmt_vtt_ts(seconds: int) -> str:
            h = seconds // 3600
            m = (seconds % 3600) // 60
            s = seconds % 60
            return f"{h:02d}:{m:02d}:{s:02d}.000"

        parts = [p.strip() for p in text.split(". ") if p.strip()]
        out = ["WEBVTT", ""]
        for i, sentence in enumerate(parts):
            start = i * 3
            end = (i + 1) * 3
            out.append(f"{_fmt_vtt_ts(start)} --> {_fmt_vtt_ts(end)}")
            out.append(sentence)
            out.append("")
        return "\n".join(out)
